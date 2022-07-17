"""
WORKING WITH PDF AND WORD DOCUMENTS

Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 347). No Starch Press. Kindle Edition. 
"""

# Extracting Text from PDFs

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 348). No Starch Press. Kindle Edition. 

import PyPDF2
import pyparsing

pdfFileObj = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
pageObj = pdfReader.getPage(0)
pageObj.extractText()
pdfFileObj.close()

# Decrypting PDFs

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 349). No Starch Press. Kindle Edition. 

import PyPDF2
pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))
pdfReader.isEncrypted
pdfReader.getPage(0)
pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))
pdfReader.decrypt('password')
pageObj = pdfReader.getPage(0)

# Copying Pages

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 351). No Starch Press. Kindle Edition. 

import PyPDF2

pdf1File = open('meetingminutes.pdf', 'rb')
pdf2File = open('meetingminutes2.pdf', 'rb')
pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
pdfWriter = PyPDF2.PdfFileWriter()

for pageNum in range(pdf1Reader.numPages):
    pageObj = pdf1Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)
    
for pageNum in range(pdf2Reader.numPages):
    pageObj = pdf2Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)
    
pdfOutputFile = open('combinedminutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()

# Rotating Pages

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 352). No Starch Press. Kindle Edition. 

import PyPDF2

minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
page = pdfReader.getPage(0)
page.rotateClockwise(90)
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(page)
resultsPdfFile = open('rotatedPage.pdf', 'wb')
pdfWriter.write(resultsPdfFile)
resultsPdfFile.close()
minutesFile.close()

# Overlaying Pages

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 353). No Starch Press. Kindle Edition. 

import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
minutesFirstPage = pdfReader.getPage(0)
pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(minutesFirstPage)

for pageNum in range(1, pdfReader.numPages):
    pageObj = pdfReader.getPage(pageNum)
    pdfWriter.addPage(pageObj)
    
resultPdfFile = open('watermarkedCover.pdf', 'wb')
pdfWriter.write(resultPdfFile)
minutesFile.close()
resultPdfFile.close()

# Reading Word Documents

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 359). No Starch Press. Kindle Edition. 

import docx

doc = docx.Document('demo.docx')
len(doc.paragraphs)
doc.paragraphs[0].text
doc.paragraphs[1].text
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text 
doc.paragraphs[1].runs[3].text 
doc.paragraphs[1].runs[4].text 

import readDocx

print(readDocx.getText('demo.docx'))

# Run Attributes

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 362). No Starch Press. Kindle Edition. 

import docx

doc = docx.Document('demo.docx')
doc.paragraphs[0].text 
doc.paragraphs[0].style # The exact id may be different:
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].text 
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, 
 doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text,
 doc.paragraphs[1].runs[4].text)
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
# doc.paragraphs[1].runs[4].imprint = True 
doc.save('restyled.docx')

# Writing Word Documents

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 364). No Starch Press. Kindle Edition. 

import docx 

doc = docx.Document()
doc.add_paragraph('Hello, world!')
paraObj1 = doc.add_paragraph('This is the second paragraph')
paraObj2 = doc.add_paragraph('This is another paragraph')
paraObj1.add_run(' Theis text is being added to the second paragraph')
doc.save('helloworld.docx')

doc.add_paragraph('Hello, world!', 'Title')

# Adding Headings

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 366). No Starch Press. Kindle Edition. 

doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

# Adding Line and Page Breaks

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 366). No Starch Press. Kindle Edition. 

doc = docx.Document()
doc.add_paragraph('THis is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is the second page!')
doc.save('twoPage.docx')

# Adding Pictures

# Sweigart, Al. Automate the Boring Stuff with Python, 2nd Edition (p. 367). No Starch Press. Kindle Edition. 

doc.add_picture('zophie.png', width=docx.shared.Inches(1),
                 height=docx.shared.Cm(4))

